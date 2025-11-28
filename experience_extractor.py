#!/usr/bin/env python3
"""
experience_extractor.py

Simple CLI that extracts total years of professional experience from a resume using
Google Gemini (hardcoded API key inside), and prints both decimal and "X years Y months".

Usage:
    python experience_extractor.py /path/to/resume.pdf

Dependencies:
    pip install google-genai pdfplumber python-docx
"""

import os
import sys
import re
import json
import hashlib
from datetime import datetime

# ------------------ HARD-CODED GEMINI API KEY ------------------
# Replace the string below with your real Gemini API key locally.
# DO NOT paste the real key into chat or share this file publicly.
GEMINI_API_KEY = "AIzaSyA-TdOVeGgIIks3Ffjup3F_UVMzOARJT2Q"

# ------------------ genai client ------------------
try:
    from google import genai
except Exception:
    raise RuntimeError("Missing dependency 'google-genai'. Install with: pip install google-genai")

client = genai.Client(api_key=GEMINI_API_KEY)

# ------------------ optional parsers ------------------
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx
except Exception:
    docx = None

# ------------------ text extraction helpers ------------------
def extract_text_from_pdf(path):
    if pdfplumber is None:
        raise RuntimeError("Install pdfplumber: pip install pdfplumber")
    parts = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            parts.append(p.extract_text() or "")
    raw = " ".join(parts)
    return re.sub(r'\s+', ' ', raw).strip()

def extract_text_from_docx(path):
    if docx is None:
        raise RuntimeError("Install python-docx: pip install python-docx")
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    raw = " ".join(parts)
    return re.sub(r'\s+', ' ', raw).strip()

def extract_text_from_txt(path):
    with open(path, "rb") as f:
        raw = f.read()
    text = raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else str(raw)
    return re.sub(r'\s+', ' ', text).strip()

def extract_text(path):
    p = path.lower()
    if p.endswith(".pdf"):
        return extract_text_from_pdf(path)
    if p.endswith(".docx"):
        return extract_text_from_docx(path)
    if p.endswith(".txt"):
        return extract_text_from_txt(path)
    raise ValueError("Unsupported file type. Supported: .pdf, .docx, .txt")

# ------------------ fallback heuristic ------------------
def fallback_years(text):
    vals = []
    if not text:
        return 0.0
    for m in re.finditer(r'(\d+(?:\.\d+)?)\s*(?:\+)?\s*(?:years?|yrs?)', text.lower()):
        try:
            vals.append(float(m.group(1)))
        except:
            pass
    return round(max(vals), 1) if vals else 0.0

# ------------------ Gemini LLM call ------------------
def ask_gemini_for_years(text, max_chars=12000):
    # trim to avoid huge payloads
    if len(text) > max_chars:
        text = text[:max_chars]

    today = datetime.now().strftime("%Y-%m-%d")
    prompt = f"""
Return ONLY JSON: {{"total_years": <float>}}.

Compute TOTAL professional work experience:
- Merge overlapping roles.
- Convert months into decimal years (1 decimal).
- Treat "present/current" as {today}.
- If unsure → return 0.0
- DO NOT output anything except JSON.

Resume:
\"\"\"{text}\"\"\"
""".strip()

    try:
        resp = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
        raw = resp.text if hasattr(resp, "text") else str(resp)

        # Try to find JSON object
        m = re.search(r'\{.*\}', raw, flags=re.DOTALL)
        if m:
            try:
                obj = json.loads(m.group(0))
                val = float(obj.get("total_years", 0.0))
                return round(val, 1)
            except Exception:
                pass

        # Try numeric fallback from output
        m2 = re.search(r'(\d+(?:\.\d+)?)', raw)
        if m2:
            return round(float(m2.group(1)), 1)

        # Final fallback
        return fallback_years(text)
    except Exception:
        return fallback_years(text)

# ------------------ convert decimal to years+months ------------------
def convert_decimal_to_human(decimal_years, method="round"):
    years = int(decimal_years)
    if method == "floor":
        months = int((decimal_years - years) * 12)
    else:
        months = int(round((decimal_years - years) * 12))

    if months >= 12:
        years += 1
        months -= 12

    if years == 0 and months == 0:
        return "0 years"
    if months == 0:
        return f"{years} years"
    if years == 0:
        return f"{months} months"
    return f"{years} years {months} months"

# ------------------ main ------------------
def main():
    if len(sys.argv) < 2:
        print("Usage: python experience_extractor.py <resume.pdf/docx/txt>")
        sys.exit(1)

    path = sys.argv[1]
    if not os.path.isfile(path):
        print("File not found:", path)
        sys.exit(1)

    print("📄 Extracting text...")
    text = extract_text(path)

    print("🤖 Asking Gemini to calculate experience...")
    decimal = ask_gemini_for_years(text)

    human = convert_decimal_to_human(decimal, method="round")  # change to "floor" if you prefer no rounding

    print(f"\n🎯 Decimal: {decimal}")
    print(f"✅ Total years of experience: {human}")

if __name__ == "__main__":
    main()
