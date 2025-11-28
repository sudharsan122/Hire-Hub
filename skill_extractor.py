#!/usr/bin/env python3
"""
skill_extractor_categorize.py

Extract skills using Gemini (hardcoded API key), normalize, dedupe, and categorize.

Usage:
    python skill_extractor_categorize.py /path/to/resume.pdf

Dependencies:
    pip install google-genai pdfplumber python-docx
WARNING:
    This version uses a hardcoded GEMINI API key (place your key locally).
"""

import os
import sys
import re
import json
import time
from datetime import datetime

# ------------------ HARD-CODED GEMINI API KEY (replace locally) ------------------
GEMINI_API_KEY = "AIzaSyA-TdOVeGgIIks3Ffjup3F_UVMzOARJT2Q"  # <-- replace locally, not in chat

# ------------------ GenAI client ------------------
try:
    from google import genai
except Exception:
    raise RuntimeError("Missing dependency 'google-genai'. Install with: pip install google-genai")

client = genai.Client(api_key=GEMINI_API_KEY)

# ------------------ File parsers ------------------
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx
except Exception:
    docx = None

def extract_text_from_pdf(path):
    if pdfplumber is None:
        raise RuntimeError("Install pdfplumber: pip install pdfplumber")
    parts = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            parts.append(p.extract_text() or "")
    return re.sub(r'\s+', ' ', " ".join(parts)).strip()

def extract_text_from_docx(path):
    if docx is None:
        raise RuntimeError("Install python-docx: pip install python-docx")
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    return re.sub(r'\s+', ' ', " ".join(parts)).strip()

def extract_text_from_txt(path):
    with open(path, "rb") as f:
        raw = f.read()
    return re.sub(r'\s+', ' ', raw.decode('utf-8', errors='ignore')).strip()

def extract_text(path):
    p = path.lower()
    if p.endswith(".pdf"):
        return extract_text_from_pdf(path)
    if p.endswith(".docx"):
        return extract_text_from_docx(path)
    if p.endswith(".txt"):
        return extract_text_from_txt(path)
    raise ValueError("Unsupported file type. Supported: .pdf, .docx, .txt")

# ------------------ Local fallback keyword list (expandable) ------------------
BASE_KEYWORDS = [
    # languages
    "c", "c++", "c#", "python", "java", "javascript", "typescript", "go", "rust", "ruby", "php", "scala", "kotlin", "swift", "r",
    # frontend
    "react", "angular", "vue", "next.js", "svelte", "html", "css", "sass", "tailwind",
    # backend/frameworks
    "node.js", "express", "django", "flask", "spring boot", "spring", "laravel", "asp.net",
    # db
    "sql", "postgresql", "mysql", "mongodb", "redis", "oracle", "mssql", "cassandra",
    # cloud/infra
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible", "helm",
    # ci/cd
    "jenkins", "github actions", "gitlab-ci", "circleci",
    # ml/data
    "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "xgboost", "lightgbm", "nlp", "opencv", "spacy",
    # data engineering
    "spark", "hadoop", "etl", "airflow",
    # embedded / platform specific
    "embedded linux", "yocto", "petalinux", "u-boot", "device tree", "kernel", "linux kernel", "bsp",
    "arm", "raspberry pi", "stm32", "nxp", "imx", "qualcomm",
    # buses / protocols
    "i2c", "spi", "uart", "gpio", "pcie", "usb", "ethernet", "can", "i2s",
    # drivers / OS
    "board bring-up", "firmware", "bootloader", "driver development", "kernel drivers", "device drivers",
    # tools
    "git", "gdb", "cmake", "make", "gcc", "clang", "vivado", "quartus", "jtag",
    # misc
    "linux", "bash", "shell", "systemd", "sysvinit", "excel", "tableau", "power bi", "docker-compose"
]

# ------------------ Normalization map ------------------
NORMALIZE_MAP = {
    # common variants -> canonical
    "react.js": "react",
    "reactjs": "react",
    "nodejs": "node.js",
    "node js": "node.js",
    "powerbi": "power bi",
    "u boot": "u-boot",
    "u_boot": "u-boot",
    "device-tree": "device tree",
    "embedded c": "c",
    "c plus plus": "c++",
    "cplusplus": "c++",
    "usb 3 0": "usb 3.0",
    "usb3.0": "usb 3.0",
    "wi fi": "wi-fi",
    "wi-fi": "wi-fi",
    "i 2 c": "i2c",
    "i 2 s": "i2s",
    "spi ": "spi",
    "gpio ": "gpio",
    "yocto project": "yocto",
    "petalinux sdk": "petalinux",
    "system verilog": "systemverilog",
    "qemu": "qemu",
    "devops": "devops",
    "microcontrollers": "microcontroller",
    "msp430ware": "msp430"
}

# ------------------ Categories keyword sets (lowercased canonical tokens) ------------------
LANGUAGES = {"c", "c++", "c#", "python", "java", "javascript", "typescript", "go", "rust", "ruby", "php", "scala", "kotlin", "swift", "r"}
TOOLS = {"git", "gdb", "cmake", "make", "gcc", "clang", "vivado", "quartus", "jtag", "docker", "helm", "ansible"}
PROTOCOLS = {"i2c", "spi", "uart", "gpio", "pcie", "usb", "ethernet", "can", "i2s", "wi-fi", "wifi", "lte", "bluetooth"}
PLATFORMS = {"embedded linux", "yocto", "petalinux", "u-boot", "raspberry pi", "stm32", "arm", "nxp", "imx", "xilinx zynq", "xilinx rfsoc", "xilinx mpsoc"}
DRIVERS = {"kernel drivers", "device drivers", "driver development", "bootloader", "board bring-up", "bsp", "firmware", "kernel", "linux kernel"}
OTHER_HINTS = {"linux", "bash", "shell", "systemd", "sysvinit", "excel", "tableau", "power bi", "etl", "spark", "hadoop"}

# ------------------ Helpers: normalize, dedupe, categorize ------------------
def normalize_token(tok: str) -> str:
    if not tok:
        return tok
    s = tok.strip().lower()
    s = re.sub(r'[\_\t]+', ' ', s)
    s = re.sub(r'\s*[\/\\]+\s*', ' / ', s)
    s = s.replace(',', ' ')
    s = re.sub(r'\s+', ' ', s)
    # apply direct map replacements for multiword patterns first
    for k, v in NORMALIZE_MAP.items():
        if s == k or s.find(k) != -1:
            s = s.replace(k, v)
    # remove stray periods unless part of version like "3.0"
    s = re.sub(r'(?<!\d)\.(?!\d)', ' ', s)
    s = s.strip()
    # common cleanup: "usb 3 0" -> "usb 3.0"
    s = re.sub(r'\busb\s+3\s+0\b', 'usb 3.0', s)
    # collapse multiple spaces
    s = re.sub(r'\s+', ' ', s)
    # final canonical tweaks
    if s == 'node':
        s = 'node.js'
    if s == 'reactjs':
        s = 'react'
    return s

def dedupe_preserve_order(items):
    seen = set()
    out = []
    for x in items:
        if x not in seen:
            seen.add(x)
            out.append(x)
    return out

def categorize_skill(skill: str):
    low = skill.lower()
    # exact matches first
    if low in LANGUAGES:
        return "languages"
    if low in TOOLS:
        return "tools"
    if low in PROTOCOLS:
        return "protocols"
    if low in PLATFORMS:
        return "platforms"
    if low in DRIVERS:
        return "drivers"
    if low in OTHER_HINTS:
        return "other"
    # heuristics: contains keywords
    if any(k in low for k in ("driver", "kernel", "bootloader", "bsp", "board bring-up", "firmware")):
        return "drivers"
    if any(k in low for k in ("linux", "embedded", "yocto", "petalinux", "u-boot")):
        return "platforms"
    if any(k in low for k in ("aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible")):
        return "tools"
    if any(k in low for k in ("i2c", "spi", "uart", "gpio", "usb", "ethernet", "can", "i2s", "bluetooth", "wi-fi")):
        return "protocols"
    if any(k in low for k in ("python", "java", "c++", "c#", "javascript", "typescript", "go", "rust")):
        return "languages"
    # fallback
    return "other"

# ------------------ Gemini prompt & call (skills) ------------------
def build_skills_prompt(resume_text, max_chars=15000):
    if len(resume_text) > max_chars:
        resume_text = resume_text[:max_chars]
    prompt = f"""
You are an extractor. Given the resume text below, return ONLY a single JSON object:

{{"skills": [<list of canonical short skill strings>] }}

Rules:
- Return skill tokens like "python", "c++", "embedded linux", "device tree", "u-boot", "yocto", "i2c", "spi", "git".
- Normalize common variants (react.js -> react, node js -> node.js, powerbi -> power bi).
- Deduplicate and return only skills actually mentioned in the resume.
- Do NOT include company names, addresses, or long descriptive sentences.
- Output EXACTLY one JSON object and nothing else.

Resume:
\"\"\"{resume_text}\"\"\"
""".strip()
    return prompt

def call_gemini_for_skills(resume_text, max_retries=2):
    prompt = build_skills_prompt(resume_text)
    for attempt in range(max_retries + 1):
        try:
            resp = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
            raw = resp.text if hasattr(resp, "text") else str(resp)
            m = re.search(r'\{.*\}', raw, flags=re.DOTALL)
            if m:
                try:
                    obj = json.loads(m.group(0))
                    skills = obj.get("skills", [])
                    if isinstance(skills, list):
                        return skills, raw
                except Exception:
                    pass
            # try array fallback
            arr_match = re.search(r'\[\s*([^\]]+?)\s*\]', raw, flags=re.DOTALL)
            if arr_match:
                items = re.findall(r'["\']([^"\']+)["\']', arr_match.group(0))
                if items:
                    return items, raw
            # otherwise fallback
            return None, raw
        except Exception:
            if attempt < max_retries:
                time.sleep(0.6*(attempt+1))
                continue
            return None, "<error>"

# ------------------ Main processing ------------------
def process_resume(path):
    text = extract_text(path)

    skills_raw, raw_model = call_gemini_for_skills(text)
    if skills_raw is None:
        # fallback local keyword scanner
        found = []
        text_low = text.lower()
        for kw in BASE_KEYWORDS:
            if re.search(r'\b' + re.escape(kw) + r'\b', text_low):
                found.append(kw)
        skills_raw = found

    # normalize tokens
    normalized = []
    for s in skills_raw:
        if not isinstance(s, str):
            continue
        tok = normalize_token(s)
        if not tok:
            continue
        normalized.append(tok)

    # dedupe preserve order
    normalized = dedupe_preserve_order(normalized)

    # final canonical fixes: apply normalization map fully
    final = []
    for s in normalized:
        s2 = s.strip()
        # map again for safety
        for k, v in NORMALIZE_MAP.items():
            if s2 == k or k in s2:
                s2 = s2.replace(k, v)
        s2 = re.sub(r'\s+', ' ', s2).strip()
        final.append(s2)

    final = dedupe_preserve_order(final)

    # categorize
    categories = {"languages": [], "tools": [], "protocols": [], "platforms": [], "drivers": [], "other": []}
    for s in final:
        cat = categorize_skill(s)
        categories[cat].append(s)

    return {
        "all_skills": final,
        "categories": categories,
        "raw_model_output_preview": (raw_model or "")[:1000]
    }

# ------------------ CLI ------------------
def main():
    if len(sys.argv) < 2:
        print("Usage: python skill_extractor_categorize.py /path/to/resume")
        sys.exit(1)

    path = sys.argv[1]
    if not os.path.isfile(path):
        print("File not found:", path)
        sys.exit(1)

    out = process_resume(path)

    print("\n✅ Cleaned skills (count={}):".format(len(out["all_skills"])))
    for s in out["all_skills"]:
        print("-", s)

    print("\n--- Categories ---")
    for cat in ["languages", "tools", "protocols", "platforms", "drivers", "other"]:
        items = out["categories"].get(cat, [])
        print(f"\n{cat.upper()} ({len(items)}):")
        if items:
            for it in items:
                print(" -", it)
        else:
            print(" - (none)")

    # optional: print small preview of model output for debugging
    # print("\n--- RAW MODEL OUTPUT PREVIEW ---")
    # print(out['raw_model_output_preview'])

if __name__ == "__main__":
    main()
